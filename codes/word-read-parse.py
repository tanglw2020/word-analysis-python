# -*- coding:utf-8 -*-

from docx import Document
from docx.shared import Inches
import time

t1 = time.perf_counter()
# document = Document('../data/2018年计算机基础及应用Ⅰ教案.docx')
# document = Document('../data/demo.docx')
document = Document('../测试.docx')
t2 = time.perf_counter()

all_paras  = document.paragraphs
print('all_paras', len(all_paras))
t3 = time.perf_counter()

# for para in all_paras:
#     # pass
#     print(para.text)
#     print()
#     print("-------")

print("style:", )
all_tables = document.tables
print('all_tables:', len(all_tables))
all_rows = all_tables[0].rows
print("rows:", len(all_rows))
all_columns = all_tables[0].columns
print("columns:", len(all_columns))
print("style:", all_tables[0].style)
all_tables = document.tables
# for row in all_rows:
#     print(row.cells(0,0))

# for table in all_tables:
#     for i in range(6):
#         print(table.cell(i,0).text)

# single_para = document.paragraphs[-1]
# for run in single_para.runs:
#     print(run)
t4 = time.perf_counter()

# styles =  document.styles
# for i, style in enumerate(styles):
#     pass
#     # print(i, style)
#     # print('__________')
# t5 = time.perf_counter()

print('timing:', t2-t1, t3-t2, t4-t3)