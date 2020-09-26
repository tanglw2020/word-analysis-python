# -*- coding:utf-8 -*-

from docx import Document
from docx.shared import Inches
import time

t1 = time.perf_counter()
# document = Document('../data/2018年计算机基础及应用Ⅰ教案.docx')
# document = Document('../data/demo.docx')
document = Document('./测试.docx')
t2 = time.perf_counter()

all_paras  = document.paragraphs
print('all_paras', len(all_paras))
t3 = time.perf_counter()

for para in all_paras:
    # pass
    print(para.text)
    # print(para.paragraph_format.tab_stops)
    # for tabstop in para.paragraph_format.tab_stops:
        # print(tabstop.leader)

    print()
    print("-------")


# single_para = document.paragraphs[-1]
# for run in single_para.runs:
#     print(run)
t4 = time.perf_counter()

# styles =  document.styles
# for i, style in enumerate(styles):
#     pass
#     # print(i, style)
#     # print('__________')
# t5 = time.perf_counter()

print('timing:', t2-t1, t3-t2, t4-t3)